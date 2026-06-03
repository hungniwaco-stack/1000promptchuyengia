from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE_DIR = ROOT / "1000 Prompt"
OUTPUT_DIR = BASE_DIR / "FORMAT_FINAL_CLEAN"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
PROMPTS_PER_SECTION = 10


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def clean_title(text: str) -> str:
    match = re.search(r"PACK\s+\d+.*", text, flags=re.IGNORECASE)
    return match.group(0).strip() if match else text.strip()


def is_toc_heading(text: str) -> bool:
    return text.strip().upper() == "MỤC LỤC"


def is_section_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s+.+\(10 prompt\)\s*$", text.strip(), flags=re.IGNORECASE))


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def apply_section_layout(target_doc, sample_doc) -> None:
    sample_section = sample_doc.sections[0]
    for section in target_doc.sections:
        section.top_margin = sample_section.top_margin
        section.bottom_margin = sample_section.bottom_margin
        section.left_margin = sample_section.left_margin
        section.right_margin = sample_section.right_margin
        section.page_width = sample_section.page_width
        section.page_height = sample_section.page_height

        section_properties = section._sectPr
        line_numbering = section_properties.find(qn("w:lnNumType"))
        if line_numbering is not None:
            section_properties.remove(line_numbering)


def set_paragraph_base(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering_properties = paragraph_properties.find(qn("w:numPr"))
    if numbering_properties is not None:
        paragraph_properties.remove(numbering_properties)

    paragraph.alignment = None
    paragraph.paragraph_format.space_before = None
    paragraph.paragraph_format.space_after = None
    paragraph.paragraph_format.line_spacing = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.right_indent = None
    paragraph.paragraph_format.tab_stops.clear_all()

    for run in paragraph.runs:
        set_run_font(run)


def set_run_font(run, *, bold=None, italic=None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = FONT_SIZE
    run.bold = bold
    run.italic = italic


def set_all_runs(paragraph, *, bold=None, italic=None) -> None:
    for run in paragraph.runs:
        set_run_font(run, bold=bold, italic=italic)


def format_prompt_paragraph(paragraph, text: str, prompt_no: int) -> None:
    prompt_prefix = "10." if prompt_no == PROMPTS_PER_SECTION else f"{prompt_no}.  "
    prompt_text = re.sub(r"^\s*(?:\d+\s*[\.\)]\s*)?", prompt_prefix, text)
    set_paragraph_text(paragraph, prompt_text)
    paragraph.paragraph_format.left_indent = Pt(31)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.tab_stops.clear_all()
    set_all_runs(paragraph, bold=None, italic=None)


def find_sample_file() -> Path:
    samples = [
        path
        for path in BASE_DIR.glob("*.docx")
        if path.name.lower().startswith("m") and "PACK 9" in path.name
    ]
    if not samples:
        raise FileNotFoundError("Missing sample file matching: M* PACK 9.docx")
    return samples[0]


def input_files() -> list[Path]:
    files = sorted(path for path in BASE_DIR.glob("PACK *.docx") if path.is_file())
    if len(files) != 10:
        raise FileNotFoundError(f"Expected 10 pack files in {BASE_DIR}, found {len(files)}")
    return files


def format_document(path: Path, sample_doc) -> Path:
    doc = Document(path)
    apply_section_layout(doc, sample_doc)

    non_empty_index = 0
    seen_toc = False
    in_toc = False
    seen_first_section = False
    prompt_no = 0

    for paragraph in list(doc.paragraphs):
        text = normalize_text(paragraph.text)
        if not text:
            remove_paragraph(paragraph)
            continue

        set_paragraph_base(paragraph)

        if non_empty_index == 0:
            set_paragraph_text(paragraph, clean_title(text))
            set_all_runs(paragraph, bold=True, italic=None)
            non_empty_index += 1
            continue

        if is_toc_heading(text):
            set_paragraph_text(paragraph, text)
            seen_toc = True
            in_toc = True
            set_all_runs(paragraph, bold=True, italic=None)
        elif seen_toc and not seen_first_section and is_section_heading(text):
            set_paragraph_text(paragraph, text)
            in_toc = False
            seen_first_section = True
            prompt_no = 0
            set_all_runs(paragraph, bold=True, italic=None)
        elif in_toc:
            set_paragraph_text(paragraph, text)
            set_all_runs(paragraph, bold=None, italic=True)
        elif is_section_heading(text):
            set_paragraph_text(paragraph, text)
            seen_first_section = True
            prompt_no = 0
            set_all_runs(paragraph, bold=True, italic=None)
        else:
            if seen_first_section:
                prompt_no = (prompt_no % PROMPTS_PER_SECTION) + 1
                format_prompt_paragraph(paragraph, text, prompt_no)
            else:
                set_paragraph_text(paragraph, text)
                set_all_runs(paragraph, bold=None, italic=None)

        non_empty_index += 1

    output_path = OUTPUT_DIR / path.name
    doc.save(output_path)
    return output_path


def main() -> None:
    sample_file = find_sample_file()
    files = input_files()

    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    sample_doc = Document(sample_file)
    print("Sample file detected")
    print(f"Formatting {len(files)} files")
    for file in files:
        output = format_document(file, sample_doc)
        print(f"OK: pack {len(list(OUTPUT_DIR.glob('*.docx')))} -> {output.parent.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
